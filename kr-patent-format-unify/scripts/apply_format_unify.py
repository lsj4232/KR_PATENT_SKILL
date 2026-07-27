#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""한국 특허 명세서 .docx에서 신설 단락 서식을 인접 본문 단락에 동기화.

사용: python apply_format_unify.py <docx_path> [--patterns "p1|p2"] [--auto-detect] [--dry-run] [--no-backup]
"""
import os, sys, re, argparse, shutil
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

# default 검출 도입어 (정규식)
DEFAULT_PATTERNS = [
    r"^본 발명의 일 실시예에 있어서,",
    r"^본 발명에 따른 .{2,40} 방법은,",
    r"^본 발명에 따른 .{2,40} 시스템은,",
    r"^도\s*\d{1,2}\s*은[^.]{3,200}이다\.",
    r"^도\s*\d{1,2}\s*는[^.]{3,200}이다\.",
    r"^이상에서 살펴본 바와 같이,\s*본 발명에 따르면,",
]

def get_full_text(para):
    return "".join(r.text for r in para.runs)

def is_header(text):
    """【...】 단독 단락이면 헤더"""
    t = text.strip()
    return bool(re.match(r"^【[^】]+】\s*$", t))

def matches_any(text, patterns):
    t = text.strip()
    for pat in patterns:
        if re.match(pat, t):
            return True
    return False

def get_rpr_signature(rpr):
    """rPr XML element를 string으로 정규화 (비교용)"""
    if rpr is None:
        return ""
    s = rpr.xml if hasattr(rpr, 'xml') else str(rpr)
    return str(s)  # XmlString → str (hashable 보장)

def find_ref_para(doc, target_idx, patterns, skip_indices):
    """target_idx 근방에서 본문 참조 단락 찾기"""
    paras = doc.paragraphs
    # 1) 직전 1~5
    for back in range(1, 6):
        i = target_idx - back
        if i < 0: break
        if i in skip_indices: continue
        p = paras[i]
        t = get_full_text(p).strip()
        if not t: continue
        if is_header(t): continue
        if matches_any(t, patterns): continue
        return i, p
    # 2) 직후 1~5
    for fwd in range(1, 6):
        i = target_idx + fwd
        if i >= len(paras): break
        if i in skip_indices: continue
        p = paras[i]
        t = get_full_text(p).strip()
        if not t: continue
        if is_header(t): continue
        if matches_any(t, patterns): continue
        return i, p
    return None, None

def copy_format(target_p, ref_p):
    """ref_p의 pPr (outlineLvl 제외) + 첫 run rPr를 target_p에 deepcopy 복사"""
    # 1) pPr 복사 — outlineLvl 제거
    src_pPr = ref_p._element.find(qn('w:pPr'))
    if src_pPr is not None:
        old_pPr = target_p._element.find(qn('w:pPr'))
        if old_pPr is not None:
            target_p._element.remove(old_pPr)
        new_pPr = deepcopy(src_pPr)
        for lvl in new_pPr.findall(qn('w:outlineLvl')):
            new_pPr.remove(lvl)
        target_p._element.insert(0, new_pPr)
    # 2) rPr 복사
    if ref_p.runs:
        ref_rPr = ref_p.runs[0]._element.find(qn('w:rPr'))
        if ref_rPr is not None:
            for run in target_p.runs:
                old_rPr = run._element.find(qn('w:rPr'))
                if old_rPr is not None:
                    run._element.remove(old_rPr)
                new_rPr = deepcopy(ref_rPr)
                run._element.insert(0, new_rPr)

def find_section_boundaries(doc):
    """본문 영역 (앞표지 ~ 【부호의 설명】 직전 또는 【청구범위】 직전) 인덱스 반환"""
    paras = doc.paragraphs
    body_end = len(paras)
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t.startswith("【부호의 설명】") or t.startswith("【청구범위】"):
            body_end = i
            break
    return 0, body_end

def extract_claims(doc):
    paras = doc.paragraphs
    cs = next((i for i, p in enumerate(paras) if "【청구범위】" in p.text), None)
    ae = next((i for i, p in enumerate(paras) if "【요약서】" in p.text), len(paras))
    if cs is None: return ""
    return "\n".join(p.text for p in paras[cs:ae])

def auto_detect_mismatches(doc, body_start, body_end, headers_skip):
    """font 불일치 자동 감지 — 직전 본문 단락과 첫 run rPr 다른 단락 찾기"""
    paras = doc.paragraphs
    suspects = []
    # 본문 평균 rPr 추출 (출현 빈도가 가장 높은 rPr signature)
    sigs = {}
    for i in range(body_start, body_end):
        p = paras[i]
        t = get_full_text(p).strip()
        if not t or is_header(t): continue
        if p.runs:
            sig = get_rpr_signature(p.runs[0]._element.find(qn('w:rPr')))
            sigs[sig] = sigs.get(sig, 0) + 1
    if not sigs: return suspects
    common_sig = max(sigs.keys(), key=lambda k: sigs[k])
    for i in range(body_start, body_end):
        p = paras[i]
        t = get_full_text(p).strip()
        if not t or is_header(t): continue
        if p.runs:
            sig = get_rpr_signature(p.runs[0]._element.find(qn('w:rPr')))
            if sig != common_sig:
                suspects.append(i)
    return suspects

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx_path")
    ap.add_argument("--patterns", default="", help="추가 도입어 정규식 (| 구분)")
    ap.add_argument("--auto-detect", action="store_true")
    ap.add_argument("--dry-run", action="store_true")
    ap.add_argument("--no-backup", action="store_true")
    args = ap.parse_args()

    src = args.docx_path
    if not os.path.exists(src):
        print(f"❌ 파일 없음: {src}", file=sys.stderr)
        sys.exit(1)

    # 패턴 구성
    patterns = list(DEFAULT_PATTERNS)
    if args.patterns:
        patterns.extend(args.patterns.split("|"))

    # 백업
    if not args.no_backup:
        base, ext = os.path.splitext(src)
        bak = base + "_서식통일전" + ext
        i = 1
        while os.path.exists(bak):
            bak = base + f"_서식통일전({i})" + ext
            i += 1
        shutil.copy2(src, bak)
        print(f"Backup: {bak}\n")

    doc = Document(src)
    paras = doc.paragraphs
    body_start, body_end = find_section_boundaries(doc)
    claims_before = extract_claims(doc)

    # 검출 대상 단락 인덱스 수집
    target_indices = []
    for i in range(body_start, body_end):
        t = get_full_text(paras[i]).strip()
        if not t: continue
        if is_header(t): continue
        if matches_any(t, patterns):
            target_indices.append(i)

    if args.auto_detect:
        extra = auto_detect_mismatches(doc, body_start, body_end, set(target_indices))
        # 도입어로 안 잡힌 것만 추가
        for i in extra:
            if i not in target_indices:
                target_indices.append(i)
        target_indices.sort()

    print(f"서식 동기화 대상 단락: {len(target_indices)}개")

    if args.dry_run:
        for ti in target_indices:
            ref_i, ref_p = find_ref_para(doc, ti, patterns, set(target_indices))
            ref_preview = (get_full_text(ref_p)[:60] + "...") if ref_p else "(없음)"
            print(f"  [{ti}] ref=[{ref_i}]: {get_full_text(paras[ti])[:80]}")
            print(f"        → ref text: {ref_preview}")
        print(f"\n(dry-run) 실제 수정 안 함.")
        return

    # 실제 적용
    applied = 0
    skipped = 0
    skip_set = set(target_indices)
    for ti in target_indices:
        ref_i, ref_p = find_ref_para(doc, ti, patterns, skip_set)
        if ref_p is None:
            skipped += 1
            print(f"  [{ti}] ⚠️ 참조 단락 못 찾음 — 스킵")
            continue
        copy_format(paras[ti], ref_p)
        applied += 1
        print(f"  [{ti}] ref=[{ref_i}]: {get_full_text(paras[ti])[:60]}...")

    print(f"\n적용: {applied}건, 스킵: {skipped}건")

    # 청구범위 무결성 검증
    claims_after = extract_claims(doc)
    if claims_after != claims_before:
        print("\n❌❌❌ 청구범위 byte 불일치 — 저장 중단", file=sys.stderr)
        sys.exit(2)
    print(f"청구범위 byte 일치: ✅")

    doc.save(src)
    print(f"Wrote: {src}")

if __name__ == "__main__":
    main()
